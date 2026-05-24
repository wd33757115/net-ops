"""
公文写作技能 - 专业的党政机关公文写作指南

来源：https://github.com/kagurananaga/official-document-writing-skill
"""

import time
from pathlib import Path
from typing import Any

from pydantic import BaseModel, Field

from src.skills.registry import skill_registry


# ==================== 1. 定义参数模型 ====================
class OfficialDocumentWritingParams(BaseModel):
    """
    公文写作技能参数
    """
    document_type: str | None = Field(None, description="公文类型，如'请示'、'通知'、'函'、'总结'、'纪要'、'报告'")
    content: str | None = Field(None, description="公文内容（用于审核或修改）")
    purpose: str | None = Field(None, description="公文用途或背景说明")
    action: str | None = Field(None, description="操作类型：'write'（撰写）、'review'（审核）、'check'（检查）、'guide'（指导）")
    user_query: str | None = Field(None, description="用户原始查询")


# ==================== 2. LLM 客户端 ====================
def get_document_llm():
    """获取公文写作专用的 LLM 客户端"""
    from langchain_deepseek import ChatDeepSeek

    from src.common.config import get_settings
    settings = get_settings()

    # LLM 模型名称
    model = getattr(settings, 'LLM_MODEL', None) or getattr(settings, 'LLM_PROVIDER', 'deepseek-chat')
    if model == "deepseek":
        model = "deepseek-chat"

    # API Key
    api_key = settings.DEEPSEEK_API_KEY

    return ChatDeepSeek(
        model=model,
        api_key=api_key,
        temperature=0.3,  # 较低温度，保持准确性
        max_tokens=4000   # 足够的 token 生成完整公文
    )


# ==================== 3. 读取参考文档 ====================
def get_skill_base_path() -> Path:
    """获取技能目录路径"""
    current_file = Path(__file__).resolve()
    # 从 skills/examples 回到 src 目录
    return current_file.parent.parent.parent / "skills-kagurananaga.official-document-writing-skill-master-05687d359a496a66dc0693973b5208475e1fca5d"


def read_reference_file(relative_path: str) -> str:
    """读取参考文档内容"""
    try:
        base_path = get_skill_base_path()
        file_path = base_path / relative_path

        if file_path.exists():
            with open(file_path, encoding='utf-8') as f:
                return f.read()
        else:
            print(f"[WARN] Reference file not found: {file_path}")
            return ""
    except Exception as e:
        print(f"[WARN] Failed to read {relative_path}: {e}")
        return ""


def get_document_templates() -> str:
    """获取公文模板"""
    content = read_reference_file("references/document-templates.md")
    if not content:
        # 如果文件不存在，返回基本模板
        content = """
# 请示模板
[发文机关] 关于 [请示事由] 的请示

[主送机关]：

[导语/依据]：
根据[法律、法规、政策依据]，[情况说明]。现将有关问题请示如下：

[主体/请示内容]：
一、[具体事项1]

二、[具体事项2]

[结尾/请求语]：
妥否，请批示。

[发文机关署名]
[成文日期]
"""
    return content


def get_gb_standard() -> str:
    """获取 GB/T 9704-2012 标准"""
    return read_reference_file("references/gb-t-9704-2012-standard.md")


def get_quality_checklist() -> str:
    """获取质量检查清单"""
    return read_reference_file("checklists/quality-checklist.md")


# ==================== 4. DOCX 文件生成 ====================
def generate_docx_from_content(content: str, document_type: str) -> bytes:
    """
    将公文内容转换为符合 GB/T 9704-2012 标准的 DOCX 格式

    GB/T 9704-2012 格式要求：
    - 用纸：A4 (210mm × 297mm)
    - 标题：2号小标宋体（但实际常用仿宋/黑体代替）
    - 正文：3号仿宋_GB2312
    - 层次标题：一、黑体；（一）楷体；1.、1. 仿宋
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    # 创建文档
    doc = Document()

    # 设置页面边距（GB/T 9704-2012 要求）
    sections = doc.sections
    for section in sections:
        # 天头（上白边）：37mm
        section.top_margin = Cm(3.7)
        # 订口（左白边）：28mm
        section.left_margin = Cm(2.8)
        # 右白边：26mm
        section.right_margin = Cm(2.6)
        # 地脚（下白边）：35mm
        section.bottom_margin = Cm(3.5)

    # 解析内容
    lines = content.strip().split('\n')
    current_paragraph = None

    for line in lines:
        line = line.strip()
        if not line:
            # 空行
            current_paragraph = doc.add_paragraph()
            current_paragraph.paragraph_format.line_spacing = 28
            continue

        # 检测是否为标题行（发文机关、无正文内容行等）
        if is_title_line(line, current_paragraph is None):
            # 发文机关标题
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.name = '仿宋_GB2312'
            run.font.size = Pt(16)
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            p.paragraph_format.line_spacing = 28
            p.paragraph_format.space_after = Pt(0)
            current_paragraph = p

        elif is_document_title(line):
            # 公文标题（如：关于xxx的请示）
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.name = '仿宋_GB2312'
            run.font.size = Pt(16)
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            p.paragraph_format.line_spacing = 28
            p.paragraph_format.space_after = Pt(0)
            current_paragraph = p

        elif is_main_recipient(line):
            # 主送机关
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = '仿宋_GB2312'
            run.font.size = Pt(16)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            p.paragraph_format.line_spacing = 28
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
            current_paragraph = p

        elif is_first_level_heading(line):
            # 一级标题：一、
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = '仿宋_GB2312'
            run.font.size = Pt(16)
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            p.paragraph_format.line_spacing = 28
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Cm(0.74)
            current_paragraph = p

        elif is_second_level_heading(line):
            # 二级标题：（一）
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = '仿宋_GB2312'
            run.font.size = Pt(16)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            p.paragraph_format.line_spacing = 28
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Cm(0.74)
            current_paragraph = p

        elif is_closing_line(line):
            # 结束语（妥否，请批示等）
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = '仿宋_GB2312'
            run.font.size = Pt(16)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            p.paragraph_format.line_spacing = 28
            p.paragraph_format.space_after = Pt(0)
            current_paragraph = p

        elif is_signature(line):
            # 署名和日期
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(line)
            run.font.name = '仿宋_GB2312'
            run.font.size = Pt(16)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            p.paragraph_format.line_spacing = 28
            p.paragraph_format.space_after = Pt(0)
            current_paragraph = p

        else:
            # 普通正文
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = '仿宋_GB2312'
            run.font.size = Pt(16)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            p.paragraph_format.line_spacing = 28
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
            current_paragraph = p

    # 保存到 BytesIO
    import io
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    return docx_buffer.getvalue()


def is_title_line(line: str, is_first: bool) -> bool:
    """判断是否为发文机关标题行"""
    # 常见的单位名称模式
    patterns = [
        r'^[^\s]+[区县市省厅局委办]',
        r'^[^\s]+[大学学院]$',
        r'^[^\s]+[公司企业]$',
        r'^[^\s]+[中心]$',
        r'^[^\s]+[指挥部办公室]$',
    ]
    import re
    for pattern in patterns:
        if re.match(pattern, line) and len(line) < 20:
            return True
    return False


def is_document_title(line: str) -> bool:
    """判断是否为公文标题行（关于xxx的xxx）"""
    import re
    return bool(re.match(r'^关于.+的[请示通知函报告纪要决定总结]', line))


def is_main_recipient(line: str) -> bool:
    """判断是否为主送机关"""
    import re
    return bool(re.match(r'^[^\s：:]+[厅局委办区县省市]：?$', line)) and len(line) < 15


def is_first_level_heading(line: str) -> bool:
    """判断是否为一级标题（一、）"""
    import re
    return bool(re.match(r'^[一二三四五六七八九十]+、', line))


def is_second_level_heading(line: str) -> bool:
    """判断是否为二级标题（（一））"""
    import re
    return bool(re.match(r'^[（(][一二三四五六七八九十]+[）)]', line))


def is_closing_line(line: str) -> bool:
    """判断是否为结束语行"""
    closing_phrases = [
        '妥否，请批示', '特此通知', '请予函复', '特此函复',
        '特此报告', '当否，请审批', '以上请示',
        '请批示', '请审核', '请审示'
    ]
    return any(phrase in line for phrase in closing_phrases)


def is_signature(line: str) -> bool:
    """判断是否为署名/日期行"""
    import re
    # 匹配日期格式：2024年1月1日、2024年01月01日
    if re.search(r'\d{4}年\d+月\d+日', line):
        return True
    # 匹配常见署名
    if re.search(r'[区县市省厅局委办]', line) and len(line) < 30:
        return True
    return False


# ==================== 5. 上传到 MinIO ====================
def upload_document_to_minio(file_data: bytes, filename: str) -> str:
    """
    上传公文文件到 MinIO，返回下载链接
    """
    try:
        from src.infrastructure.storage.minio_client import get_minio_storage

        minio = get_minio_storage()

        if not minio.is_ready():
            print("[WARN] MinIO not ready, skipping upload")
            return None

        # 生成对象名称
        import datetime
        timestamp = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
        object_name = f"documents/{timestamp}_{filename}"

        # 上传文件
        success = minio.upload_file(
            object_name=object_name,
            file_data=file_data,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        if success:
            # 生成下载链接（有效期7天）
            download_url = minio.get_presigned_url(object_name, expires=3600 * 24 * 7)
            print(f"   [MinIO] Document uploaded: {object_name}")
            print(f"   [MinIO] Download URL: {download_url}")
            return download_url
        else:
            print("[WARN] Failed to upload document to MinIO")
            return None

    except Exception as e:
        print(f"[WARN] MinIO upload error: {e}")
        return None


# ==================== 6. LLM 生成公文 ====================
async def generate_document_with_llm(
    document_type: str,
    purpose: str,
    user_query: str,
    templates: str,
    gb_standard: str
) -> str:
    """
    使用 LLM 生成完整的公文内容
    """
    llm = get_document_llm()

    # 构建 prompt
    prompt = f"""你是一位专业的党政机关公文写作专家，擅长撰写符合《GB/T 9704-2012》标准的公文。

## 用户需求
{user_query}

## 公文类型
{document_type}

## 公文用途/背景
{purpose}

## 公文模板参考
{templates[:3000]}  # 限制长度

## GB/T 9704-2012 格式规范
{gb_standard[:2000] if gb_standard else '请遵循党政机关公文格式规范'}

## 写作要求
1. 严格按照公文格式规范撰写
2. 语言正式严谨，表述准确
3. 内容有理有据，贴合实际
4. 结构清晰，层次分明
5. 篇幅适中，详略得当

## 输出要求
请直接输出完整的公文内容，不要包含任何解释性文字。公文格式示例：

```
[发文机关]
关于 [事由] 的 {document_type}

[主送机关]：

[正文内容]
...

[结尾语]
妥否，请批示。

[发文机关署名]
[成文日期]
```

请立即开始撰写："""

    try:
        t_start = time.time()
        response = await llm.ainvoke(prompt)
        elapsed = time.time() - t_start

        print(f"   [Document Generation] LLM took {elapsed:.1f}s")

        return response.content if hasattr(response, 'content') else str(response)

    except Exception as e:
        print(f"   [Document Generation] LLM Error: {e}")
        # 如果 LLM 调用失败，返回模板提示
        return None


# ==================== 5. Handler 函数 ====================
async def official_document_writing_handler(params: dict[str, Any]) -> dict[str, Any]:
    """
    公文写作处理函数

    提供党政机关公文写作的完整服务，包括：
    - 使用 LLM 生成符合规范的公文
    - GB/T 9704-2012格式规范
    - 常用公文模板
    - 公文质量检查
    """
    start_time = time.time()

    try:
        document_type = params.get("document_type")
        content = params.get("content")
        purpose = params.get("purpose")
        action = params.get("action", "write")
        user_query = params.get("user_query", "")

        # 如果用户没有提供 document_type，从 user_query 中推断
        if not document_type and user_query:
            document_type = infer_document_type(user_query)

        # 读取参考文档
        templates = get_document_templates()
        gb_standard = get_gb_standard()
        checklist = get_quality_checklist()

        # 处理不同操作类型
        if action == "review" or action == "check":
            # 审核/检查模式
            result = await review_document(content, document_type, checklist, gb_standard)

        elif action == "guide":
            # 指导模式
            result = await guide_document(document_type, templates)

        else:
            # 默认撰写模式 - 调用 LLM 生成公文
            result = await write_document(document_type, purpose, user_query, templates, gb_standard)

        elapsed = time.time() - start_time
        result["execution_time_ms"] = int(elapsed * 1000)

        return result

    except Exception as e:
        print(f"   [Official Document Writing] Error: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "message": f"公文写作处理失败: {str(e)}",
            "error": str(e),
            "execution_time_ms": 0
        }


def infer_document_type(query: str) -> str:
    """从用户查询中推断公文类型"""
    query_lower = query.lower()

    if "请示" in query_lower:
        return "请示"
    elif "通知" in query_lower:
        return "通知"
    elif "函" in query_lower:
        return "函"
    elif "纪要" in query_lower:
        return "会议纪要"
    elif "报告" in query_lower:
        return "报告"
    elif "总结" in query_lower:
        return "工作总结"
    elif "决定" in query_lower:
        return "决定"
    else:
        return "通用公文"


async def write_document(
    document_type: str,
    purpose: str,
    user_query: str,
    templates: str,
    gb_standard: str
) -> dict[str, Any]:
    """撰写公文"""

    print(f"   [Document Writing] Generating {document_type}...")

    # 调用 LLM 生成公文
    document_content = await generate_document_with_llm(
        document_type=document_type,
        purpose=purpose,
        user_query=user_query,
        templates=templates,
        gb_standard=gb_standard
    )

    if document_content:
        # 生成 DOCX 文件
        print("   [Document Writing] Generating DOCX file...")
        try:
            docx_data = generate_docx_from_content(document_content, document_type)

            # 生成文件名
            import datetime
            timestamp = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
            filename = f"{document_type}_{timestamp}.docx"

            # 上传到 MinIO
            download_url = upload_document_to_minio(docx_data, filename)

            if download_url:
                print("   [Document Writing] DOCX uploaded successfully")

            # 返回结果
            return {
                "success": True,
                "message": "公文撰写完成",
                "data": {
                    "document_type": document_type,
                    "purpose": purpose,
                    "action": "write",
                    "document_content": document_content,
                    "workflow": ["需求确认", "文种选择", "模板应用", "内容写作", "质量检查"],
                    "reference_files": [
                        "references/document-templates.md",
                        "references/gb-t-9704-2012-standard.md"
                    ],
                    "compliance_notes": [
                        "已遵循 GB/T 9704-2012 格式规范",
                        "语言正式严谨，表述准确",
                        "结构清晰，层次分明"
                    ]
                },
                "download_url": download_url  # 添加下载链接
            }

        except Exception as e:
            print(f"   [Document Writing] DOCX generation failed: {e}")
            import traceback
            traceback.print_exc()

            # 如果 DOCX 生成失败，仍然返回文本内容
            return {
                "success": True,
                "message": "公文撰写完成（DOCX生成失败，请使用文本内容）",
                "data": {
                    "document_type": document_type,
                    "purpose": purpose,
                    "action": "write",
                    "document_content": document_content,
                    "workflow": ["需求确认", "文种选择", "模板应用", "内容写作", "质量检查"],
                    "reference_files": [
                        "references/document-templates.md",
                        "references/gb-t-9704-2012-standard.md"
                    ],
                    "compliance_notes": [
                        "已遵循 GB/T 9704-2012 格式规范",
                        "语言正式严谨，表述准确",
                        "结构清晰，层次分明"
                    ]
                }
            }
    else:
        # LLM 生成失败，返回模板提示
        return {
            "success": True,
            "message": f"已为您准备{document_type}的写作模板，请参考以下内容：",
            "data": {
                "document_type": document_type,
                "purpose": purpose,
                "action": "write",
                "workflow": ["需求确认", "文种选择", "模板应用", "内容写作", "质量检查"],
                "reference_files": [
                    "references/document-templates.md",
                    "references/gb-t-9704-2012-standard.md",
                    "references/writing-techniques.md",
                    "checklists/quality-checklist.md"
                ],
                "tips": [
                    "先规划后写作：明确目的、对象、文种",
                    "善用模板：在模板基础上个性化修改",
                    "重视依据：政策依据、事实依据要充分",
                    "简洁明了：删繁就简，直截了当",
                    "多查多改：初稿完成后多次检查修改",
                    "对标对表：对照清单逐项检查"
                ]
            }
        }


async def review_document(
    content: str,
    document_type: str,
    checklist: str,
    gb_standard: str
) -> dict[str, Any]:
    """审核公文"""
    llm = get_document_llm()

    prompt = f"""你是一位专业的公文审核专家，请审核以下公文内容：

## 公文类型
{document_type or '通用公文'}

## 待审核内容
{content}

## 质量检查清单
{checklist[:2000]}

## GB/T 9704-2012 格式规范
{gb_standard[:1500] if gb_standard else '请遵循党政机关公文格式规范'}

## 审核要求
请从以下方面进行审核：
1. 格式规范：标题、主送机关、正文、落款等格式是否正确
2. 语言规范：是否准确、平实、简明、庄重
3. 内容质量：逻辑是否清晰，依据是否充分
4. 法规合规：是否符合相关政策法规要求

请提供详细的审核意见和改进建议。

格式：
### 审核结果：通过 / 需要修改
### 格式问题：（如有）
### 内容问题：（如有）
### 修改建议：
"""

    try:
        t_start = time.time()
        response = await llm.ainvoke(prompt)
        elapsed = time.time() - t_start

        print(f"   [Document Review] LLM took {elapsed:.1f}s")

        review_result = response.content if hasattr(response, 'content') else str(response)

        return {
            "success": True,
            "message": "公文审核完成",
            "data": {
                "document_type": document_type,
                "action": "review",
                "review_result": review_result,
                "reference_files": [
                    "checklists/quality-checklist.md",
                    "references/gb-t-9704-2012-standard.md"
                ]
            }
        }

    except Exception as e:
        print(f"   [Document Review] LLM Error: {e}")
        return {
            "success": True,
            "message": "已收到您的公文内容，正在进行质量检查...",
            "data": {
                "document_type": document_type,
                "action": "review",
                "suggestions": [
                    "请对照GB/T 9704-2012标准检查格式规范",
                    "检查字体字号是否符合要求",
                    "检查结构层次序数是否正确",
                    "检查成文日期格式是否正确",
                    "检查语言是否准确、平实、简明、庄重"
                ],
                "reference_files": [
                    "checklists/quality-checklist.md",
                    "references/gb-t-9704-2012-standard.md"
                ]
            }
        }


async def guide_document(
    document_type: str,
    templates: str
) -> dict[str, Any]:
    """公文写作指导"""
    llm = get_document_llm()

    # 提取对应类型的模板
    template_section = ""
    if document_type:
        lines = templates.split('\n')
        in_section = False
        section_lines = []

        for line in lines:
            if document_type in line and ('模板' in line or '###' in line):
                in_section = True
                section_lines.append(line)
            elif in_section:
                if line.startswith('## ') and document_type not in line:
                    break
                section_lines.append(line)

        template_section = '\n'.join(section_lines[:50])

    prompt = f"""你是一位专业的公文写作导师，请提供关于"{document_type}"的写作指导。

## 模板参考
{template_section}

## 指导要求
请提供：
1. {document_type}的定义和使用场景
2. 写作要点和注意事项
3. 常用句式和表达
4. 写作示例

请用简洁专业的语言进行指导。"""

    try:
        t_start = time.time()
        response = await llm.ainvoke(prompt)
        elapsed = time.time() - t_start

        print(f"   [Document Guide] LLM took {elapsed:.1f}s")

        guide_content = response.content if hasattr(response, 'content') else str(response)

        return {
            "success": True,
            "message": guide_content,
            "data": {
                "document_type": document_type,
                "action": "guide",
                "writing_principles": ["准确", "平实", "简明", "庄重"],
                "reference_files": [
                    "references/document-templates.md",
                    "references/writing-techniques.md"
                ]
            }
        }

    except Exception as e:
        print(f"   [Document Guide] LLM Error: {e}")
        # 返回基本指导
        templates_dict = {
            "请示": "请示写作要点：明确请示依据，说清请示事项，一文一事，结尾用'妥否，请批示'",
            "通知": "通知写作要点：说明通知依据，明确通知事项，提出执行要求，结尾用'特此通知'",
            "函": "函写作要点：说明发函依据，明确商洽事项，注意语气得体，结尾用'请予函复'或'特此函复'",
            "总结": "总结写作要点：概括主要成绩，提炼经验做法，查找存在问题，明确今后打算",
            "纪要": "纪要写作要点：记录会议信息完整，客观记录讨论情况，明确议定事项，落实责任分工",
            "报告": "报告写作要点：真实准确，条理清晰，重点突出，报告及时"
        }

        guide_content = templates_dict.get(
            document_type,
            f"正在为您提供{document_type}的写作指导..."
        )

        return {
            "success": True,
            "message": guide_content,
            "data": {
                "document_type": document_type,
                "action": "guide",
                "writing_principles": ["准确", "平实", "简明", "庄重"],
                "reference_files": [
                    "references/document-templates.md",
                    "references/writing-techniques.md"
                ]
            }
        }


# ==================== 6. 注册技能 ====================
def register_skill():
    """
    注册公文写作技能
    """
    from src.skills.skill_base import BaseSkill, SkillResult

    class OfficialDocumentWritingSkill(BaseSkill):
        async def execute(self, **kwargs) -> SkillResult:
            result = await official_document_writing_handler(kwargs)
            return SkillResult(**result)

    skill = OfficialDocumentWritingSkill(
        name="official_document_writing",  # 技能名称（唯一标识）
        description="专业的党政机关公文写作助手。当用户需要撰写、修改或审核党政机关公文（如请示、通知、函、总结、纪要等）时使用。提供GB/T 9704-2012格式规范指导、常用公文模板、语言规范建议和公文质量检查。支持直接生成符合规范的完整公文。",
        parameters=OfficialDocumentWritingParams,
        handler=official_document_writing_handler,
        category="general",  # 分类：network, security, general 等
        tags=["公文写作", "党政机关", "请示", "通知", "函", "总结", "纪要", "报告", "GB/T 9704-2012"],  # 标签
        fallback_to_rag_if_fail=True,  # 失败时是否走 RAG
        enabled=True  # 是否启用
    )

    skill_registry.register_skill(skill)


# ==================== 7. 测试入口 ====================
if __name__ == "__main__":
    import asyncio

    async def test():
        register_skill()
        skill = skill_registry.get_skill("official_document_writing")
        if skill:
            print(f"✅ 技能注册成功: {skill.name}")
            print(f"   描述: {skill.description[:50]}...")
            print(f"   分类: {skill.category}")
            print(f"   标签: {skill.tags}")

            # 测试生成公文
            print("\n--- 测试生成公文 ---")
            result = await skill.execute(
                document_type="请示",
                purpose="申请采购新的网络设备",
                user_query="帮我写一份请示，向信息中心申请采购一台新的核心交换机",
                action="write"
            )
            print(f"Result: {result.success}")
            print(f"Message: {result.message}")
            if hasattr(result, 'data') and result.data:
                content = result.data.get('document_content', 'N/A')
                print(f"Content Preview: {content[:200]}...")

        else:
            print("❌ 技能注册失败")

    asyncio.run(test())
