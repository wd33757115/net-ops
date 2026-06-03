# SPDX-FileCopyrightText: 2026 wangdong <wangdong5919@163.com>
# SPDX-License-Identifier: Apache-2.0

"""公文 DOCX 渲染：结构化 JSON → Word（优先 python-docx 直排）。"""

from __future__ import annotations

import logging
from io import BytesIO
from pathlib import Path

from src.skills.official_document.docx_format import (
    add_styled_paragraph,
    setup_a4_margins,
)
from src.skills.official_document.schema import OfficialDocumentJSON

logger = logging.getLogger(__name__)

SKILL_ASSETS = Path(__file__).resolve().parent.parent / "official-document-writing" / "assets" / "templates"
TEMPLATE_VERSION = "2"  # 行距修复：v1 模板 line_spacing=28 导致一行一页

TEMPLATE_MAP = {
    "请示": "请示.docx",
    "通知": "通知.docx",
    "函": "函.docx",
    "报告": "报告.docx",
    "工作总结": "报告.docx",
    "总结": "报告.docx",
    "会议纪要": "会议纪要.docx",
    "纪要": "会议纪要.docx",
}


def get_skill_assets_dir() -> Path:
    return SKILL_ASSETS


def ensure_default_templates() -> None:
    """生成/升级 docxtpl 占位模板（行距已修正）。"""
    assets = get_skill_assets_dir()
    assets.mkdir(parents=True, exist_ok=True)
    version_file = assets / ".template_version"

    if version_file.exists() and version_file.read_text(encoding="utf-8").strip() == TEMPLATE_VERSION:
        return

    for old in assets.glob("*.docx"):
        old.unlink(missing_ok=True)

    template_specs = {
        "default.docx": _default_placeholder_lines(),
        **{name: _default_placeholder_lines() for name in TEMPLATE_MAP.values()},
    }
    for filename, lines in template_specs.items():
        _write_placeholder_docx(assets / filename, lines)

    version_file.write_text(TEMPLATE_VERSION, encoding="utf-8")
    logger.info("已重建公文模板 v%s", TEMPLATE_VERSION)


def _default_placeholder_lines() -> list[str]:
    return [
        "{{ issuer }}",
        "{{ title }}",
        "{{ main_recipient }}：",
        "{{ main_body_text }}",
        "{{ closing }}",
        "{{ signature_org }}",
        "{{ signature_date }}",
    ]


def _write_placeholder_docx(path: Path, lines: list[str]) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    from src.skills.official_document.docx_format import apply_fixed_line_spacing, style_run

    doc = Document()
    setup_a4_margins(doc)

    for idx, line in enumerate(lines):
        p = doc.add_paragraph()
        if idx <= 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_fixed_line_spacing(p)
        style_run(p.add_run(line))
        if idx == 2:
            p.paragraph_format.first_line_indent = Cm(0)

    doc.save(str(path))


def resolve_template_path(doc_type: str) -> Path:
    ensure_default_templates()
    filename = TEMPLATE_MAP.get(doc_type, "default.docx")
    path = get_skill_assets_dir() / filename
    if not path.exists():
        path = get_skill_assets_dir() / "default.docx"
    return path


def render_document_bytes(document: OfficialDocumentJSON | dict) -> bytes:
    """根据结构化 JSON 渲染 DOCX。"""
    if isinstance(document, dict):
        document = OfficialDocumentJSON.model_validate(document)

    # 直排更可控；docxtpl 仅作备选
    try:
        return _render_structured_bytes(document)
    except Exception as exc:
        logger.warning("结构化直排失败，尝试 docxtpl: %s", exc)
        return _render_docxtpl_bytes(document)


def _render_structured_bytes(document: OfficialDocumentJSON) -> bytes:
    """按 JSON 字段逐段排版（推荐路径）。"""
    from docx import Document

    doc = Document()
    setup_a4_margins(doc)

    if document.issuer:
        add_styled_paragraph(doc, document.issuer, center=True, bold=True)
    add_styled_paragraph(doc, document.title, center=True, bold=True)
    add_styled_paragraph(doc, f"{document.main_recipient}：")

    if document.main_body.opening:
        add_styled_paragraph(doc, document.main_body.opening, first_line_indent=0.74)

    for section in document.main_body.sections:
        if section.heading:
            add_styled_paragraph(doc, section.heading, first_line_indent=0.74, bold=True)
        if section.content:
            add_styled_paragraph(doc, section.content, first_line_indent=0.74)

    if document.main_body.closing:
        add_styled_paragraph(doc, document.main_body.closing)

    add_styled_paragraph(doc, document.signature.org, right=True)
    add_styled_paragraph(doc, document.signature.date, right=True)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _render_docxtpl_bytes(document: OfficialDocumentJSON) -> bytes:
    context = document.to_render_context()
    template_path = resolve_template_path(document.doc_type)

    try:
        from docxtpl import DocxTemplate
    except ImportError:
        logger.warning("docxtpl 未安装，使用结构化直排")
        return _render_structured_bytes(document)

    tpl = DocxTemplate(str(template_path))
    tpl.render(context)
    buffer = BytesIO()
    tpl.save(buffer)
    return buffer.getvalue()
